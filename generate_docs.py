import os
from docx import Document
from docx.shared import Pt, Inches, RGBColor
from docx.enum.text import WD_ALIGN_PARAGRAPH

def add_title(doc, text):
    title = doc.add_heading(text, 0)
    title.alignment = WD_ALIGN_PARAGRAPH.CENTER

def add_heading(doc, text, level=1):
    doc.add_heading(text, level=level)

def add_paragraph(doc, text, bold=False):
    p = doc.add_paragraph()
    run = p.add_run(text)
    if bold:
        run.bold = True
    return p

def add_bullet(doc, text):
    doc.add_paragraph(text, style='List Bullet')

def add_placeholder(doc, text):
    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    run = p.add_run(f"\n[ 🖼️ ILAGAY DITO ANG PICTURE: {text} ]\n")
    run.font.color.rgb = RGBColor(255, 0, 0)
    run.bold = True

def create_system_doc(folder):
    doc = Document()
    add_title(doc, "FormRequest System - Overall Documentation")
    
    add_heading(doc, "1. Ano ang FormRequest System?")
    add_paragraph(doc, "Ang FormRequest System ay ginawa para mapadali, mapabilis, at maging digital ang buong proseso ng pag-request at pag-approve ng Gatepass (personnel, vehicle, at materials). Imbes na gumamit ng papel na madalas nawawala o matagal mapirmahan, lahat ngayon ay online na. Kasama rin dito ang pag-monitor ng guard, pag-schedule ng mga sasakyan ng HRAD (PAS), at pagpapadali sa pag-approve ng mga boss gamit ang digital signature.")
    
    add_heading(doc, "2. Summary ng mga Natapos na Phases")
    add_bullet(doc, "Phase 1 - Phase 3: Paggawa ng pinaka-foundation ng system. Kasama dito ang login system, ang basic na pag-request ng Vehicle at Personnel gatepass, at pag-setup ng database na naka-link sa system ninyo.")
    add_bullet(doc, "Phase 4 - Phase 6: Dito idinagdag ang Material Gatepass para sa mga gamit na inilalabas. Sinama rin ang digital approval system kung saan makakapag-pirma (e-signature) ang mga approvers diretso sa system.")
    add_bullet(doc, "Phase 7 - Phase 9: Pinasok natin ang QR Code scanning para sa mga Guards para isang scan na lang, recorded na agad ang IN at OUT. Inayos din ang pagpapakita ng live PDF form bago i-print.")
    add_bullet(doc, "Phase 10 - Phase 12: Inayos natin ang scheduling calendar system. Ginawang mas matibay ang security, mas malinaw ang pag-assign ng mga sasakyan (PAS), at in-optimize ang view para sa mga naka-mobile phone.")
    add_bullet(doc, "Phase 13+: Final adjustments sa UI, gaya ng pag-aayos ng layout pag pinrint ang A6 form, pag-aayos ng bugs sa calendar, at pagpapaayos ng transition ng mga pages para smooth gamitin.")

    add_heading(doc, "3. Future Improvements (Mga Pwedeng Idagdag)")
    add_bullet(doc, "Iba't-ibang uri ng Forms: Pwede natin idagdag ang mga IT Helpdesk Requests, HR Leave Forms, o Maintenance Work Orders sa iisang system na lang para di na hiwa-hiwalay.")
    add_bullet(doc, "Reporting at Analytics Dashboard: Maganda kung may graph na magpapakita kung ilang gatepass ang na-approve per month, sino ang pinaka-madalas lumabas, at anong sasakyan ang pinaka-gamit para sa madaling audit.")
    add_bullet(doc, "Email at SMS Notifications: Para automatic na makaka-receive ng text o email ang nag-request kapag na-approve na ng boss ang kanyang gatepass.")

    doc.save(os.path.join(folder, "1_System_Overview.docx"))

def create_requester_manual(folder):
    doc = Document()
    add_title(doc, "User Manual: Para sa mga Nagre-request (Requester)")
    
    add_heading(doc, "Paano Gumawa ng Gatepass Request?")
    add_paragraph(doc, "Bilang normal na user, ikaw ang kadalasang nagagawa ng request para makalabas ng kumpanya o makagamit ng sasakyan.")
    add_bullet(doc, "1. Mag-login gamit ang iyong Employee ID at Password.")
    add_placeholder(doc, "Login Screen")
    add_bullet(doc, "2. Sa main dashboard, pumili kung anong klase ng gatepass ang kailangan mo (Personnel, Vehicle, o Material).")
    add_placeholder(doc, "Dashboard na may Buttons ng Gatepass")
    add_bullet(doc, "3. Punan ang mga detalye tulad ng pangalan, date, purpose, at destination. Siguraduhing tama ang impormasyon.")
    add_bullet(doc, "4. I-click ang 'Submit'. Hintayin na ma-approve ito ng inyong superior o ng taong in-charge.")
    
    add_heading(doc, "Paano I-check ang Status at I-print?")
    add_bullet(doc, "1. Pumunta sa 'My Requests' tab.")
    add_placeholder(doc, "My Requests List")
    add_bullet(doc, "2. Makikita mo dito kung 'Pending', 'Approved', o 'Rejected' na ang iyong request.")
    add_bullet(doc, "3. Kung Approved na, magkakaroon ng 'Print' button. I-click ito para lumabas ang pormal na Gatepass form na may QR Code at pirma.")
    add_placeholder(doc, "Printable Form View")
    
    doc.save(os.path.join(folder, "2_User_Manual_Requester.docx"))

def create_approver_manual(folder):
    doc = Document()
    add_title(doc, "User Manual: Para sa mga Approvers (Superiors & President)")
    
    add_heading(doc, "Paano Mag-Approve o Mag-Reject ng Request?")
    add_paragraph(doc, "Bilang approver, ikaw ang may huling say kung papayagan ang request ng isang empleyado.")
    add_bullet(doc, "1. Mag-login sa system at pumunta sa 'Approval Dashboard'.")
    add_placeholder(doc, "Approval Dashboard showing Pending Requests")
    add_bullet(doc, "2. I-click ang 'Review' button sa tabi ng pangalan ng nag-request.")
    add_bullet(doc, "3. Basahin ang detalye (Saan pupunta, ano ang purpose).")
    add_placeholder(doc, "Document Review Modal")
    add_bullet(doc, "4. Para mag-approve, ilagay ang iyong e-signature sa digital canvas na nasa screen. Pwede kang pumirma gamit ang mouse o touchscreen.")
    add_placeholder(doc, "Signature Pad")
    add_bullet(doc, "5. I-click ang 'Approve'. Kung hindi naman pwede, i-click ang 'Reject' at maglagay ng dahilan kung bakit.")
    
    doc.save(os.path.join(folder, "3_User_Manual_Approver.docx"))

def create_pas_manual(folder):
    doc = Document()
    add_title(doc, "User Manual: Para sa PAS / HRAD (Vehicle Management)")
    
    add_heading(doc, "Paano I-manage ang Sasakyan at Driver?")
    add_paragraph(doc, "Kayo ang nagko-kontrol kung sino ang sasakay sa aling truck, sino ang driver, at anong schedule.")
    add_bullet(doc, "1. Mag-login at buksan ang 'Vehicle Schedule' o Calendar view.")
    add_placeholder(doc, "Calendar View Dashboard")
    add_bullet(doc, "2. Kapag may pumasok na Vehicle Request na kailangan ng sasakyan ng kumpanya, pupunta ito sa inyong listahan bilang 'Pending Schedule'.")
    add_placeholder(doc, "Pending Schedules List")
    add_bullet(doc, "3. I-click ang request, pagkatapos ay pumili ng available na sasakyan at driver sa dropdown menu.")
    add_bullet(doc, "4. I-save ito para ma-update ang kalendaryo. Makikita na ito ng lahat na occupied ang sasakyang iyon sa araw na iyon.")
    
    add_heading(doc, "Fixed / Permanent Schedules")
    add_bullet(doc, "1. Para sa mga araw-araw na byahe, i-click ang 'Manage Fixed Schedules'.")
    add_placeholder(doc, "Manage Fixed Schedules Modal")
    add_bullet(doc, "2. Dito pwede kayong mag-set ng byahe every Monday, Tuesday, etc., para hindi na kailangang i-type araw-araw.")
    
    doc.save(os.path.join(folder, "4_User_Manual_PAS_HRAD.docx"))

def create_guard_manual(folder):
    doc = Document()
    add_title(doc, "User Manual: Para sa mga Security Guard")
    
    add_heading(doc, "Paano Mag-Scan ng Gatepass?")
    add_paragraph(doc, "Trabaho ng guard na siguraduhing authorized ang mga taong lumalabas at pumapasok.")
    add_bullet(doc, "1. Buksan ang Guard Scanner page sa tablet o computer.")
    add_placeholder(doc, "Guard Scanner Page")
    add_bullet(doc, "2. Itapat ang printed Gatepass na may QR Code sa camera.")
    add_bullet(doc, "3. Pag na-scan, lalabas sa screen ang impormasyon ng gatepass. Kulay GREEN kung valid at Approved, kulay RED kung Rejected o bawal.")
    add_placeholder(doc, "Scan Result (Valid / Invalid)")
    add_bullet(doc, "4. I-click ang 'Mark as OUT' kapag lumabas na ang sasakyan o tao.")
    add_bullet(doc, "5. I-click ang 'Mark as IN' kapag bumalik na sila. Automatic itong mase-save sa system kaya hindi na kailangan mag-sulat sa logbook.")
    
    doc.save(os.path.join(folder, "5_User_Manual_Guard.docx"))

def create_admin_manual(folder):
    doc = Document()
    add_title(doc, "User Manual: Para sa System Administrator")
    
    add_heading(doc, "Paano I-manage ang System at Users?")
    add_paragraph(doc, "Bilang Admin, ikaw ang may kakayahang magdagdag ng user, mag-reset ng password, at mag-maintain ng records.")
    add_bullet(doc, "1. Pumunta sa 'Admin Panel'.")
    add_placeholder(doc, "Admin Panel Dashboard")
    add_bullet(doc, "2. Sa User Management, pwede kang mag-add ng bagong empleyado o mag-deactivate ng nag-resign na.")
    add_placeholder(doc, "User Management Page")
    add_bullet(doc, "3. Sa Master Lists, dito mo ia-update ang listahan ng mga Trucks, Drivers, at Destinations para palaging bago ang pagpipilian ng mga users sa form.")
    add_placeholder(doc, "Master List Settings")
    add_bullet(doc, "4. Makikita mo rin ang buong Audit Trail kung mayroong nagka-aberya o gustong i-trace kung sino ang nag-approve ng request.")
    
    doc.save(os.path.join(folder, "6_User_Manual_Admin.docx"))

if __name__ == "__main__":
    desktop = os.path.join(os.environ['USERPROFILE'], 'Desktop')
    folder = os.path.join(desktop, 'FormRequest_Documentation')
    
    if not os.path.exists(folder):
        os.makedirs(folder)
        
    create_system_doc(folder)
    create_requester_manual(folder)
    create_approver_manual(folder)
    create_pas_manual(folder)
    create_guard_manual(folder)
    create_admin_manual(folder)
    
    print(f"Docs generated successfully in {folder}")
